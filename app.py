from flask import Flask, render_template, request
import os
import io
from spellchecker import SpellChecker
import pandas as pd
import re
from docx import Document
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 🔥 SET PATH TESSERACT
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# LOAD EXCEL
data_pegawai = pd.read_excel("data_pegawai.xlsx")

# BERSIHKAN NAMA
def bersihkan_nama(teks):
    teks = teks.lower()
    teks = re.sub(r'[^a-z\s]', '', teks)
    return teks.strip()

# 🔥 DOCX NORMAL
def extract_docx_text(file):
    doc = Document(io.BytesIO(file.read()))
    texts = []

    for p in doc.paragraphs:
        texts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)

    return "\n".join(texts)

# 🔥 OCR (UNTUK TEXTBOX)
def extract_with_ocr(file_bytes):
    images = convert_from_bytes(file_bytes)
    text = ""

    for img in images:
        text += pytesseract.image_to_string(img)

    return text


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/cek", methods=["POST"])
def cek():
    file = request.files['file']

    if file.filename == '':
        return "Tidak ada file dipilih"

    file_bytes = file.read()

    # 🔥 1. COBA BACA NORMAL
    try:
        full_text = extract_docx_text(io.BytesIO(file_bytes))
    except:
        full_text = ""

    # 🔥 2. TAMBAH OCR (BIAR TEXTBOX KEAMBIL)
    try:
        ocr_text = extract_with_ocr(file_bytes)
        full_text += "\n" + ocr_text
    except:
        pass

    text_lower = full_text.lower()
    text_bersih = bersihkan_nama(full_text)

    # SPELL CHECK
    spell = SpellChecker()
    words = full_text.split()
    salah = spell.unknown(words)

    # VALIDASI JABATAN
    hasil_validasi = []

    for i in range(len(data_pegawai)):
        nama_asli = str(data_pegawai.loc[i, 'nama'])
        jabatan_asli = str(data_pegawai.loc[i, 'jabatan_aktif'])

        nama_excel = bersihkan_nama(nama_asli)
        jabatan_excel = jabatan_asli.lower().strip()

        if nama_excel in text_bersih:
            if jabatan_excel in text_lower:
                hasil_validasi.append(f"{nama_asli} ✅ sesuai ({jabatan_asli})")
            else:
                hasil_validasi.append(f"{nama_asli} ❌ jabatan salah! Seharusnya: {jabatan_asli}")

    # STATUS
    if len(salah) > 0 or any("❌" in v for v in hasil_validasi):
        status = "Perlu Revisi ❌"
    else:
        status = "Layak Cetak ✅"

    return render_template(
        "hasil.html",
        isi=full_text,
        salah=salah,
        status=status,
        validasi=hasil_validasi
    )


if __name__ == "__main__":
    app.run(debug=True)